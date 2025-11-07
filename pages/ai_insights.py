import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import base64
import io
import uuid
from services.data_security import DataSecurity

# Document processing imports
try:
    import pytesseract
    from PIL import Image
    import PyPDF2
    import docx
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False
    st.warning("Document processing libraries not installed. Install with: pip install pytesseract pillow PyPDF2 python-docx")

# AI processing (mock implementation - replace with actual AI service)
try:
    import openai
    AI_SERVICE_AVAILABLE = True
except ImportError:
    AI_SERVICE_AVAILABLE = False

def show():
    # Professional header styling
    st.markdown("""
    <style>
    /* Match main app background */
    .stApp {
        background: linear-gradient(135deg, 
            #1a0b2e 0%,
            #2d1b4e 15%,
            #1e3a8a 35%,
            #0f172a 50%,
            #1e3a8a 65%,
            #16537e 85%,
            #0891b2 100%) !important;
        min-height: 100vh;
        position: relative;
    }
    
    /* Geometric overlay pattern */
    .stApp::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: 
            radial-gradient(circle at 20% 30%, rgba(168, 85, 247, 0.1) 0%, transparent 50%),
            radial-gradient(circle at 80% 70%, rgba(14, 165, 233, 0.1) 0%, transparent 50%),
            radial-gradient(circle at 50% 50%, rgba(59, 130, 246, 0.05) 0%, transparent 50%);
        pointer-events: none;
    }
    /* Sidebar styling - must be in each page file */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #667eea 0%, #764ba2 100%) !important;
        padding: 0 !important;
    }
    
    [data-testid="stSidebar"] > div:first-child {
        padding: 2rem 1rem !important;
    }
    
    [data-testid="stSidebar"] .css-17eq0hr {
        color: white !important;
    }
    
    [data-testid="stSidebar"] label {
        color: white !important;
        font-weight: 600 !important;
    }
    
    [data-testid="stSidebar"] button {
        background: rgba(255,255,255,0.15) !important;
        color: white !important;
        border: 1px solid rgba(255,255,255,0.2) !important;
        border-radius: 12px !important;
        padding: 0.75rem 1rem !important;
        width: 100% !important;
    }
    
    [data-testid="stSidebar"] button:hover {
        background: rgba(255,255,255,0.25) !important;
    }
    .ai-header {
        background: rgba(30, 58, 138, 0.6);
        backdrop-filter: blur(10px);
        padding: 3rem 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
        border: 1px solid rgba(59, 130, 246, 0.2);
    }
    .ai-header h1 {
        color: white;
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
        font-weight: 700;
    }
    /* Default: Light text everywhere EXCEPT inside white containers */
    body, .stApp, p, span, div {
        color: #e2e8f0;
    }
    
    /* Headers - light */
    h1, h2, h3, h4, h5, h6 {
        color: #f1f5f9 !important;
    }
    
    /* Dark text ONLY inside white expander boxes */
    [data-testid="stExpander"] [data-testid="stExpanderDetails"] {
        background: white !important;
    }
    
    [data-testid="stExpander"] [data-testid="stExpanderDetails"] * {
        color: #1e293b !important;
    }
    
    /* ADD THESE NEW RULES: */
    
    /* Query Suggestions - white boxes with dark text */
    .element-container:has(.stButton) {
        background: white !important;
        border-radius: 8px !important;
        padding: 0.5rem !important;
    }
    
    .element-container:has(.stButton) * {
        color: #1e293b !important;
    }
    
    /* Any white/light background containers should have dark text */
    [style*="background: white"],
    [style*="background: #fff"],
    [style*="background-color: white"],
    [style*="background-color: #fff"],
    [style*="background: rgb(255, 255, 255)"] {
        color: #1e293b !important;
    }
    
    [style*="background: white"] *,
    [style*="background: #fff"] *,
    [style*="background-color: white"] *,
    [style*="background-color: #fff"] *,
    [style*="background: rgb(255, 255, 255)"] * {
        color: #1e293b !important;
    }
    
    /* White containers in general */
    .stContainer[style*="background"],
    div[style*="background: white"],
    div[style*="background-color: white"],
    div[style*="background: #ffffff"] {
        background: white !important;
    }
    
    .stContainer[style*="background"] *,
    div[style*="background: white"] *,
    div[style*="background-color: white"] *,
    div[style*="background: #ffffff"] * {
        color: #1e293b !important;
    }
    
    /* Markdown content in white boxes */
    .stMarkdown[style*="background"] * {
        color: #1e293b !important;
    }
    
    /* Dark text in forms */
    [data-testid="stForm"] {
        background: rgba(255, 255, 255, 0.95) !important;
    }
    
    [data-testid="stForm"] * {
        color: #1e293b !important;
    }
    
    /* Metrics - keep colored */
    [data-testid="stMetricValue"] {
        color: #60a5fa !important;
    }
    
    [data-testid="stMetricLabel"] {
        color: #cbd5e1 !important;
    }
    
    /* Input fields */
    input, textarea, select {
        color: #1e293b !important;
        background: white !important;
    }
    
    /* Buttons */
    .stButton button * {
        color: white !important;
    }
    
    /* Info/warning/error boxes - keep light text */
    .stAlert, .stSuccess, .stWarning, .stError, .stInfo {
        color: #1e293b !important;
    }

    /* Dropdown menus - dark text */
    [data-baseweb="select"] [role="listbox"],
    [data-baseweb="select"] [role="option"],
    [data-baseweb="popover"] {
        background: white !important;
    }
    
    [data-baseweb="select"] [role="listbox"] *,
    [data-baseweb="select"] [role="option"] *,
    [data-baseweb="popover"] * {
        color: #1e293b !important;
    }
    
    /* Dropdown list items */
    [data-baseweb="menu"] li,
    [data-baseweb="menu"] li *,
    div[role="listbox"] li,
    div[role="listbox"] li * {
        color: #1e293b !important;
        background: white !important;
    }
    
    /* Select/dropdown text */
    .stSelectbox [data-baseweb="select"] > div {
        color: #1e293b !important;
    }
    .metric-card {
        background: rgba(30, 41, 59, 0.8);
        backdrop-filter: blur(10px);
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.2);
        border-left: 4px solid #3b82f6;
        border: 1px solid rgba(59, 130, 246, 0.2);
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 12px 24px;
        background-color: rgba(30, 41, 59, 0.6);
        backdrop-filter: blur(10px);
        border-radius: 8px;
        font-weight: 600;
        color: #cbd5e1;
        border: 1px solid rgba(59, 130, 246, 0.2);
    }
    .stTabs [aria-selected="true"] {
        background-color: rgba(59, 130, 246, 0.8);
        color: white;
    }
    

    </style>
    <div class="ai-header">
        <h1>🤖 AI Legal Insights</h1>
        <p>AI-powered document analysis, contract review, and legal intelligence</p>
    </div>
    """, unsafe_allow_html=True)
    
    DataSecurity.require_auth("AI Legal Insights")

    # Main tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📄 Document Intelligence", 
        "📊 Contract Analytics", 
        "🔍 Legal Research", 
        "⚖️ Risk Assessment", 
        "📈 Practice Analytics",
        "🔮 Predictive Insights"
    ])
    
    with tab1:
        show_document_intelligence()
    
    with tab2:
        show_contract_analytics()
    
    with tab3:
        show_legal_research()
    
    with tab4:
        show_risk_assessment()
    
    with tab5:
        show_practice_analytics()
    
    with tab6:
        show_predictive_insights()

def show_document_intelligence():
    """Document reading and analysis interface"""
    st.subheader("📄 Document Intelligence & Analysis")
    
    if not DOCUMENT_PROCESSING_AVAILABLE:
        st.error("Document processing not available. Please install required libraries.")
        st.code("pip install pytesseract pillow PyPDF2 python-docx")
        return
    
    # Document upload section
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("#### Upload Document or Image")
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'],
            help="Supported formats: PDF, Word, Text, and Image files"
        )
        
        # Analysis options
        if uploaded_file:
            st.markdown("#### Analysis Options")
            analysis_type = st.selectbox(
                "Select Analysis Type",
                [
                    "Full Document Analysis",
                    "Contract Review",
                    "Key Terms Extraction",
                    "Legal Risk Assessment",
                    "Document Summarization",
                    "Clause Analysis",
                    "Compliance Check"
                ]
            )
            
            analyze_button = st.button("🔍 Analyze Document", type="primary")
    
    with col2:
        if uploaded_file:
            st.markdown("#### File Information")
            st.write(f"**Filename:** {uploaded_file.name}")
            st.write(f"**File type:** {uploaded_file.type}")
            st.write(f"**File size:** {uploaded_file.size / 1024:.1f} KB")
            
            # Preview for images
            if uploaded_file.type.startswith('image/'):
                st.markdown("#### Preview")
                image = Image.open(uploaded_file)
                st.image(image, width=200)
    
    # Document analysis results
    if uploaded_file and analyze_button:
        with st.spinner("Analyzing document..."):
            # Extract text from document
            extracted_text = extract_document_text(uploaded_file)
            
            if extracted_text:
                # Perform AI analysis
                analysis_results = perform_document_analysis(extracted_text, analysis_type)
                
                # Display results
                display_analysis_results(analysis_results, extracted_text, uploaded_file.name)
            else:
                st.error("Could not extract text from the document. Please check the file format.")
    
    # Batch processing section
    st.divider()
    st.markdown("#### Batch Document Processing")
    
    col1, col2 = st.columns(2)
    
    with col1:
        batch_files = st.file_uploader(
            "Upload multiple documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Process multiple documents at once"
        )
        
        if batch_files:
            batch_analysis_type = st.selectbox(
                "Batch Analysis Type",
                ["Document Classification", "Contract Comparison", "Risk Scoring", "Key Terms Extraction"],
                key="batch_analysis"
            )
            
            if st.button("🔄 Process Batch"):
                process_batch_documents(batch_files, batch_analysis_type)
    
    with col2:
        if batch_files:
            st.markdown("#### Batch Summary")
            st.write(f"**Files selected:** {len(batch_files)}")
            total_size = sum(file.size for file in batch_files)
            st.write(f"**Total size:** {total_size / 1024:.1f} KB")
            
            for file in batch_files[:5]:  # Show first 5 files
                st.write(f"• {file.name}")
            
            if len(batch_files) > 5:
                st.write(f"... and {len(batch_files) - 5} more files")

def extract_document_text(uploaded_file):
    """Extract text from various document formats"""
    try:
        file_type = uploaded_file.type
        
        if file_type == "application/pdf":
            return extract_text_from_pdf(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(uploaded_file)
        elif file_type == "text/plain":
            return str(uploaded_file.read(), "utf-8")
        elif file_type.startswith('image/'):
            return extract_text_from_image(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return None

def extract_text_from_pdf(uploaded_file):
    """Extract text from PDF files"""
    try:
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if not text.strip():
            return "PDF contains no extractable text. It may be scanned or image-based."
        
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_docx(uploaded_file):
    """Extract text from Word documents"""
    try:
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        doc = docx.Document(uploaded_file)  # ✅ Fixed: Added 'Document'
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            return "Word document contains no text."
        
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        st.info("Make sure python-docx is installed: pip install python-docx")
        return None

def extract_text_from_image(uploaded_file):
    """Extract text from images using OCR"""
    try:
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        image = Image.open(uploaded_file)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang='eng')
        
        if not text.strip():
            return "No text detected in image. Image may be too low quality or contain no text."
        
        return text
    except pytesseract.TesseractNotFoundError:
        st.error("Tesseract OCR is not installed on your system")
        st.info("""
        To install Tesseract:
        - **Windows**: Download from https://github.com/UB-Mannheim/tesseract/wiki
        - **Mac**: `brew install tesseract`
        - **Linux**: `sudo apt-get install tesseract-ocr`
        """)
        return None
    except Exception as e:
        st.error(f"Error performing OCR: {str(e)}")
        return None

def perform_document_analysis(text, analysis_type):
    """Perform AI analysis on extracted text"""
    # Mock AI analysis - replace with actual AI service calls
    
    analysis_results = {
        "summary": "",
        "key_terms": [],
        "risk_score": 0,
        "compliance_issues": [],
        "recommendations": [],
        "clauses": [],
        "entities": [],
        "sentiment": "neutral"
    }
    
    if analysis_type == "Full  Analysis":
        analysis_results = perform_full_analysis(text)
    elif analysis_type == "Contract Review":
        analysis_results = perform_contract_review(text)
    elif analysis_type == "Key Terms Extraction":
        analysis_results = extract_key_terms(text)
    elif analysis_type == "Legal Risk Assessment":
        analysis_results = assess_legal_risks(text)
    elif analysis_type == "Document Summarization":
        analysis_results = summarize_document(text)
    elif analysis_type == "Clause Analysis":
        analysis_results = analyze_clauses(text)
    elif analysis_type == "Compliance Check":
        analysis_results = check_compliance(text)
    
    return analysis_results

def perform_full_analysis(text):
    """Comprehensive document analysis using actual text"""
    if not text or len(text.strip()) < 50:
        return {"summary": "Document too short for analysis", "key_terms": [], "risk_score": 0}
    
    from collections import Counter
    import re
    
    words = text.split()
    word_count = len(words)
    sentences = [s.strip() for s in text.split('.') if s.strip()]
    
    # Extract actual key terms
    word_freq = Counter(word.lower().strip('.,!?;:()[]') for word in words if len(word) > 3)
    key_terms = [word for word, count in word_freq.most_common(10)]
    
    # Extract entities
    orgs = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|LLC|Corp|Company|Ltd)\b', text)
    dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
    locations = [loc for loc in ['New York', 'California', 'Delaware', 'Texas'] if loc in text]
    
    # Calculate risk score based on actual keywords
    risk_keywords = ['liability', 'penalty', 'breach', 'indemnify', 'terminate', 'damages']
    risk_score = min(10, sum(2 for kw in risk_keywords if kw in text.lower()) / 2)
    
    # Detect document type
    doc_type = "General Document"
    if any(word in text.lower() for word in ['agreement', 'contract', 'parties']):
        doc_type = "Contract"
    elif any(word in text.lower() for word in ['complaint', 'plaintiff', 'defendant']):
        doc_type = "Legal Pleading"
    
    return {
        "summary": f"This document contains {word_count} words across {len(sentences)} sentences. Most frequent terms suggest it is a {doc_type.lower()}.",
        "key_terms": key_terms[:10],
        "risk_score": round(risk_score, 1),
        "compliance_issues": ["Review required - automated compliance check not available"],
        "recommendations": [
            f"Document contains {len(orgs)} organization references",
            f"Found {len(dates)} date references",
            "Manual legal review recommended for compliance"
        ],
        "entities": list(set(orgs[:5] + locations)),
        "sentiment": "neutral",
        "document_type": doc_type,
        "confidence": 0.75
    }

def perform_contract_review(text):
    """Analyze contract-specific elements"""
    return {
        "summary": "Contract review identifies key commercial terms and potential risk areas.",
        "clauses": [
            {"type": "Payment Terms", "status": "Present", "risk": "Low"},
            {"type": "Termination", "status": "Unclear", "risk": "Medium"},
            {"type": "Intellectual Property", "status": "Missing", "risk": "High"},
            {"type": "Liability", "status": "Present", "risk": "Low"}
        ],
        "key_terms": ["payment", "delivery", "warranty", "indemnification"],
        "risk_score": 4.1,
        "missing_clauses": ["Force Majeure", "Governing Law", "Dispute Resolution"],
        "recommendations": [
            "Add intellectual property protection clauses",
            "Clarify termination procedures",
            "Include dispute resolution mechanisms"
        ]
    }

def extract_key_terms(text):
    """Extract important legal terms and concepts from actual text"""
    if not text:
        return {"summary": "No text to analyze", "key_terms": [], "term_analysis": []}
    
    from collections import Counter
    
    # Extract all meaningful words (length > 3)
    words = [word.lower().strip('.,!?;:()[]"\'') for word in text.split() if len(word) > 3]
    word_freq = Counter(words)
    
    # Define legal keywords to prioritize
    legal_keywords = ["shall", "agreement", "party", "parties", "obligation", "right", "liability", 
                     "termination", "breach", "indemnify", "warranty", "payment", "fee",
                     "contract", "clause", "provision", "covenant", "entity", "jurisdiction"]
    
    # Build term analysis
    terms = []
    for keyword in legal_keywords:
        count = word_freq.get(keyword, 0)
        if count > 0:
            terms.append({
                "term": keyword,
                "frequency": count,
                "importance": "high" if count > 3 else "medium"
            })
    
    # Add other frequent terms not in legal keywords
    for word, count in word_freq.most_common(15):
        if word not in legal_keywords and count > 2:
            terms.append({
                "term": word,
                "frequency": count,
                "importance": "medium" if count > 5 else "low"
            })
    
    return {
        "summary": f"Extracted {len(terms)} key terms from the document. Most frequent: {terms[0]['term'] if terms else 'none'}",
        "key_terms": [term["term"] for term in terms[:10]],
        "term_analysis": sorted(terms, key=lambda x: x["frequency"], reverse=True)[:15],
        "most_frequent": max(terms, key=lambda x: x["frequency"]) if terms else None
    }

def assess_legal_risks(text):
    """Assess potential legal risks in the document"""
    return {
        "summary": "Risk assessment identifies potential legal exposures and compliance issues.",
        "risk_score": 3.7,
        "high_risk_areas": [
            "Unlimited liability exposure",
            "Vague termination clauses",
            "Missing intellectual property protection"
        ],
        "medium_risk_areas": [
            "Payment terms ambiguity",
            "Force majeure provisions"
        ],
        "low_risk_areas": [
            "Standard confidentiality clauses",
            "Basic governing law provisions"
        ],
        "recommendations": [
            "Implement liability caps",
            "Strengthen IP protection clauses",
            "Add comprehensive force majeure provisions"
        ]
    }

def summarize_document(text):
    """Generate document summary"""
    word_count = len(text.split())
    sentences = text.split('.')
    
    return {
        "summary": f"This document is approximately {word_count} words long and contains {len(sentences)} sentences. It appears to be a legal document covering contractual obligations and terms.",
        "key_points": [
            "Establishes contractual relationship between parties",
            "Defines mutual obligations and responsibilities",
            "Includes standard legal protections",
            "Specifies terms for agreement termination"
        ],
        "document_structure": {
            "word_count": word_count,
            "sentence_count": len(sentences),
            "estimated_reading_time": f"{word_count // 250} minutes"
        }
    }

def analyze_clauses(text):
    """Analyze specific contract clauses from actual text"""
    if not text:
        return {"summary": "No text to analyze", "clauses": []}
    
    text_lower = text.lower()
    
    # Define clause types and their detection keywords
    clause_checks = {
        "Payment Terms": ["payment", "fee", "compensation", "remuneration", "pay"],
        "Termination": ["termination", "terminate", "end this agreement", "cancel"],
        "Force Majeure": ["force majeure", "act of god", "beyond control"],
        "Confidentiality": ["confidential", "non-disclosure", "proprietary"],
        "Intellectual Property": ["intellectual property", "copyright", "trademark", "ip rights"],
        "Liability": ["liability", "liable", "indemnify", "indemnification"],
        "Governing Law": ["governed by", "jurisdiction", "applicable law"],
        "Dispute Resolution": ["dispute", "arbitration", "mediation", "litigation"]
    }
    
    clauses = []
    for clause_name, keywords in clause_checks.items():
        present = any(keyword in text_lower for keyword in keywords)
        keyword_count = sum(text_lower.count(kw) for kw in keywords)
        
        if present:
            quality = "Good" if keyword_count > 2 else "Fair" if keyword_count > 0 else "Weak"
            issues = [] if keyword_count > 2 else ["Limited detail - may need expansion"]
            suggestions = [] if keyword_count > 2 else [f"Consider adding more detailed {clause_name.lower()} provisions"]
        else:
            quality = "Missing"
            issues = [f"{clause_name} not found in document"]
            suggestions = [f"Add {clause_name.lower()} clause for completeness"]
        
        clauses.append({
            "name": clause_name,
            "present": present,
            "quality": quality,
            "issues": issues,
            "suggestions": suggestions
        })
    
    present_count = sum(1 for c in clauses if c["present"])
    
    return {
        "summary": f"Found {present_count} of {len(clauses)} standard clauses in the document.",
        "clauses": clauses
    }

def check_compliance(text):
    """Check document compliance with legal standards"""
    return {
        "summary": "Compliance check evaluates adherence to legal and regulatory requirements.",
        "compliance_score": 7.2,
        "compliant_areas": [
            "Privacy policy requirements",
            "Basic contract formation elements",
            "Signature requirements"
        ],
        "non_compliant_areas": [
            "Missing mandatory disclosures",
            "Unclear dispute resolution",
            "Incomplete governing law specification"
        ],
        "recommendations": [
            "Add required regulatory disclosures",
            "Specify complete governing law and jurisdiction",
            "Include mandatory consumer protection clauses"
        ]
    }

def display_analysis_results(results, extracted_text, filename):
    """Display comprehensive analysis results"""
    st.markdown(f"### Analysis Results for: {filename}")
    
    # Summary section
    with st.expander("📋 Executive Summary", expanded=True):
        st.write(results.get("summary", "No summary available"))
        
        if "risk_score" in results:
            col1, col2, col3 = st.columns(3)
            with col1:
                risk_score = results["risk_score"]
                risk_color = "red" if risk_score > 6 else "orange" if risk_score > 3 else "green"
                st.metric("Risk Score", f"{risk_score}/10", delta=None)
            
            with col2:
                if "compliance_score" in results:
                    st.metric("Compliance Score", f"{results['compliance_score']}/10")
            
            with col3:
                if "confidence" in results:
                    st.metric("AI Confidence", f"{results['confidence']*100:.1f}%")
    
    # Key findings tabs
    if any(key in results for key in ["key_terms", "clauses", "entities", "recommendations"]):
        tab1, tab2, tab3, tab4 = st.tabs(["🔑 Key Terms", "📄 Clauses", "🏢 Entities", "💡 Recommendations"])
        
        with tab1:
            if "key_terms" in results and results["key_terms"]:
                st.markdown("#### Most Important Terms")
                for term in results["key_terms"][:10]:
                    st.write(f"• **{term}**")
            else:
                st.info("No key terms extracted")
        
        with tab2:
            if "clauses" in results:
                for clause in results["clauses"]:
                    status_color = {"Present": "green", "Missing": "red", "Unclear": "orange"}.get(clause["status"], "gray")
                    risk_color = {"Low": "green", "Medium": "orange", "High": "red"}.get(clause["risk"], "gray")
                    
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        st.write(f"**{clause['type']}**")
                    with col2:
                        st.markdown(f":{status_color}[{clause['status']}]")
                    with col3:
                        st.markdown(f":{risk_color}[{clause['risk']} Risk]")
            else:
                st.info("No clause analysis available")
        
        with tab3:
            if "entities" in results and results["entities"]:
                st.markdown("#### Identified Entities")
                for entity in results["entities"]:
                    st.write(f"• {entity}")
            else:
                st.info("No entities identified")
        
        with tab4:
            if "recommendations" in results and results["recommendations"]:
                st.markdown("#### AI Recommendations")
                for i, rec in enumerate(results["recommendations"], 1):
                    st.write(f"{i}. {rec}")
            else:
                st.info("No recommendations available")
    
    # Risk assessment section
    if any(key in results for key in ["high_risk_areas", "compliance_issues"]):
        with st.expander("⚠️ Risk Assessment"):
            col1, col2 = st.columns(2)
            
            with col1:
                if "high_risk_areas" in results:
                    st.markdown("##### High Risk Areas")
                    for risk in results["high_risk_areas"]:
                        st.error(f"🔴 {risk}")
                
                if "medium_risk_areas" in results:
                    st.markdown("##### Medium Risk Areas")
                    for risk in results["medium_risk_areas"]:
                        st.warning(f"🟡 {risk}")
            
            with col2:
                if "compliance_issues" in results:
                    st.markdown("##### Compliance Issues")
                    for issue in results["compliance_issues"]:
                        st.warning(f"⚖️ {issue}")
    
    # Document text preview
    with st.expander("📖 Extracted Text Preview"):
        preview_text = extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text
        st.text_area("Document Content", preview_text, height=200, disabled=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Total characters:** {len(extracted_text):,}")
        with col2:
            st.write(f"**Word count:** {len(extracted_text.split()):,}")
    
    # Save results option
    if st.button("💾 Save Analysis Results"):
        save_analysis_results(results, filename)

def process_batch_documents(batch_files, analysis_type):
    """Process multiple documents at once"""
    results_data = []
    progress_bar = st.progress(0)
    
    for i, file in enumerate(batch_files):
        st.write(f"Processing: {file.name}")
        
        # Extract text
        text = extract_document_text(file)
        
        if text:
            # Perform analysis
            if analysis_type == "Document Classification":
                doc_type = classify_document(text)
                results_data.append({
                    "filename": file.name,
                    "type": doc_type,
                    "confidence": 0.85,
                    "word_count": len(text.split())
                })
            elif analysis_type == "Risk Scoring":
                risk_score = assess_document_risk(text)
                results_data.append({
                    "filename": file.name,
                    "risk_score": risk_score,
                    "word_count": len(text.split())
                })
            elif analysis_type == "Key Terms":
                    terms_result = extract_key_terms(text)
                    results_data.append({
                        "Filename": file.name,
                        "Terms Found": len(terms_result.get("key_terms", [])),
                        "Most Frequent": terms_result.get("most_frequent", {}).get("term", "N/A") if terms_result.get("most_frequent") else "N/A",
                        "Word Count": len(text.split()),
                        "Status": "✓ Success"
                })
                
            elif analysis_type == "Clause Analysis":
                clause_result = analyze_clauses(text)
                clauses = clause_result.get("clauses", [])
                present_clauses = sum(1 for c in clauses if c.get("present"))
                results_data.append({
                    "Filename": file.name,
                    "Clauses Found": present_clauses,
                    "Total Checked": len(clauses),
                    "Completion": f"{round(present_clauses/len(clauses)*100)}%" if clauses else "0%",
                    "Status": "✓ Success"
                })
            
            elif analysis_type == "Full Analysis":
                full_result = perform_full_analysis(text)
                results_data.append({
                    "Filename": file.name,
                    "Doc Type": full_result.get("document_type", "Unknown"),
                    "Risk Score": full_result.get("risk_score", 0),
                    "Key Terms": len(full_result.get("key_terms", [])),
                    "Word Count": len(text.split()),
                    "Status": "✓ Success"
                })
        progress_bar.progress((i + 1) / len(batch_files))
    
    # Display batch results
    if results_data:
        st.success(f"Processed {len(results_data)} documents successfully!")
        df = pd.DataFrame(results_data)
        st.dataframe(df, use_container_width=True)
        
        # Download results
        csv = df.to_csv(index=False)
        st.download_button(
            label="📥 Download Results as CSV",
            data=csv,
            file_name=f"batch_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def classify_document(text):
    """Classify document type based on content"""
    # Simple keyword-based classification
    if any(word in text.lower() for word in ["agreement", "contract", "parties"]):
        return "Contract"
    elif any(word in text.lower() for word in ["complaint", "plaintiff", "defendant"]):
        return "Legal Pleading"
    elif any(word in text.lower() for word in ["memo", "memorandum", "analysis"]):
        return "Legal Memorandum"
    else:
        return "General Legal Document"

def assess_legal_risks(text):
    """Assess potential legal risks based on actual document content"""
    if not text:
        return {"summary": "No text to analyze", "risk_score": 0}
    
    text_lower = text.lower()
    
    # Define risk indicators
    high_risk_indicators = {
        "unlimited liability": ["unlimited", "liability", "without limit"],
        "unclear termination": ["may terminate" if "may terminate" in text_lower and "notice" not in text_lower else ""],
        "missing IP protection": [] if any(term in text_lower for term in ["intellectual property", "ip rights", "copyright"]) else ["no IP protection"]
    }
    
    medium_risk_indicators = {
        "payment ambiguity": [] if "payment" in text_lower and any(t in text_lower for t in ["date", "within", "days"]) else ["payment terms unclear"],
        "force majeure": [] if "force majeure" in text_lower else ["no force majeure clause"]
    }
    
    low_risk_indicators = {
        "confidentiality": ["confidential" in text_lower or "non-disclosure" in text_lower],
        "governing law": ["governed by" in text_lower or "jurisdiction" in text_lower]
    }
    
    # Calculate actual risks
    high_risk_areas = []
    for risk, indicators in high_risk_indicators.items():
        if indicators and indicators[0]:
            high_risk_areas.append(risk.title())
    
    medium_risk_areas = []
    for risk, indicators in medium_risk_indicators.items():
        if indicators and indicators[0]:
            medium_risk_areas.append(risk.title())
    
    low_risk_areas = [area.replace("_", " ").title() for area, present in low_risk_indicators.items() if present]
    
    # Calculate risk score
    risk_keywords = ['liability', 'penalty', 'breach', 'damages', 'indemnify', 'lawsuit', 'litigation']
    risk_count = sum(3 for kw in risk_keywords if kw in text_lower)
    protection_keywords = ['limit', 'cap', 'maximum', 'insurance', 'indemnification']
    protection_count = sum(1 for kw in protection_keywords if kw in text_lower)
    risk_score = max(0, min(10, risk_count - protection_count))
    
    return {
        "summary": f"Risk assessment identified {len(high_risk_areas)} high-risk and {len(medium_risk_areas)} medium-risk areas.",
        "risk_score": round(risk_score, 1),
        "high_risk_areas": high_risk_areas if high_risk_areas else ["No critical risks detected"],
        "medium_risk_areas": medium_risk_areas if medium_risk_areas else ["No medium risks detected"],
        "low_risk_areas": low_risk_areas if low_risk_areas else ["Standard provisions present"],
        "recommendations": [
            f"Address {len(high_risk_areas)} high-priority risk areas" if high_risk_areas else "Continue monitoring for risks",
            "Consider professional legal review for comprehensive assessment",
            "Ensure all critical clauses are present and detailed"
        ]
    }

def save_analysis_results(results, filename):
    """Save analysis results to session state"""
    if 'document_analyses' not in st.session_state:
        st.session_state.document_analyses = []
    
    analysis_record = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "timestamp": datetime.now(),
        "results": results
    }
    
    st.session_state.document_analyses.append(analysis_record)
    st.success("Analysis results saved!")

def show_contract_analytics():
    """Contract-specific analytics and insights"""
    st.subheader("📊 Contract Analytics")
    
    # Mock contract analytics data
    contract_data = {
        "contract_types": ["Service Agreement", "NDA", "Employment", "Vendor", "Lease"],
        "counts": [25, 18, 12, 15, 8],
        "avg_risk_scores": [3.2, 2.1, 4.5, 3.8, 2.9]
    }
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Contract type distribution
        fig1 = px.pie(
            values=contract_data["counts"], 
            names=contract_data["contract_types"],
            title="Contract Distribution by Type"
        )
        st.plotly_chart(fig1, use_container_width=True)
    
    with col2:
        # Risk scores by contract type
        fig2 = px.bar(
            x=contract_data["contract_types"],
            y=contract_data["avg_risk_scores"],
            title="Average Risk Score by Contract Type"
        )
        fig2.update_layout(yaxis_title="Risk Score", xaxis_title="Contract Type")
        st.plotly_chart(fig2, use_container_width=True)

def show_legal_research():
    """Legal research and case law analysis"""
    st.subheader("🔍 Legal Research Assistant")
    
    research_query = st.text_input("Enter your legal research query:")
    
    if research_query:
        st.info("Legal research functionality would integrate with legal databases like Westlaw, LexisNexis, or public case law APIs")
        
        # Mock research results
        st.markdown("#### Research Results")
        results = [
            {"title": "Smith v. Johnson (2023)", "relevance": "95%", "jurisdiction": "Federal"},
            {"title": "Contract Law Principles", "relevance": "88%", "jurisdiction": "State"},
            {"title": "Recent Precedent Analysis", "relevance": "82%", "jurisdiction": "Federal"}
        ]
        
        for result in results:
            with st.expander(f"{result['title']} - {result['relevance']} relevant"):
                st.write(f"**Jurisdiction:** {result['jurisdiction']}")
                st.write("**Summary:** Mock legal research result...")

def show_risk_assessment():
    """Risk assessment dashboard"""
    st.subheader("⚖️ Legal Risk Assessment")
    
    # Risk metrics with professional styling
    col1, col2, col3, col4 = st.columns(4)
    
    metrics_html = """
    <div class="metric-card" style="text-align: center;">
        <h4 style="color: #64748b; font-size: 0.9rem; margin-bottom: 0.5rem;">{title}</h4>
        <p style="font-size: 1.8rem; font-weight: 700; color: {color}; margin: 0;">{value}</p>
        <p style="color: #64748b; font-size: 0.85rem; margin-top: 0.5rem;">{change}</p>
    </div>
    """
    
    with col1:
        st.markdown(metrics_html.format(
            title="High Risk Documents",
            value="12",
            change="+3 this month",
            color="#ef4444"
        ), unsafe_allow_html=True)
    
    with col2:
        st.markdown(metrics_html.format(
            title="Compliance Score",
            value="8.2/10",
            change="+0.5 improvement",
            color="#10b981"
        ), unsafe_allow_html=True)
    
    with col3:
        st.markdown(metrics_html.format(
            title="Average Risk Score",
            value="3.4/10",
            change="-0.2 lower",
            color="#3b82f6"
        ), unsafe_allow_html=True)
    
    with col4:
        st.markdown(metrics_html.format(
            title="Documents Reviewed",
            value="156",
            change="+23 this week",
            color="#3b82f6"
        ), unsafe_allow_html=True)

def show_practice_analytics():
    """Practice management analytics"""
    st.subheader("📈 Practice Analytics")
    
    # Document processing metrics with professional cards
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #1e3a8a; margin-bottom: 1rem;">Documents Processed</h3>
            <p style="font-size: 2rem; font-weight: 700; color: #3b82f6; margin: 0;">1,247</p>
            <p style="color: #64748b; margin-top: 0.5rem;">Processing Accuracy: 94.2%</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #1e3a8a; margin-bottom: 1rem;">Time Saved</h3>
            <p style="font-size: 2rem; font-weight: 700; color: #3b82f6; margin: 0;">142 hours</p>
            <p style="color: #64748b; margin-top: 0.5rem;">Cost Savings: $28,400</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #1e3a8a; margin-bottom: 1rem;">AI Performance</h3>
            <p style="font-size: 2rem; font-weight: 700; color: #3b82f6; margin: 0;">87.5%</p>
            <p style="color: #64748b; margin-top: 0.5rem;">User Satisfaction: 4.6/5</p>
        </div>
        """, unsafe_allow_html=True)

def show_predictive_insights():
    """Predictive analytics and insights"""
    st.subheader("🔮 Predictive Legal Insights")
    
    st.info("Predictive insights would analyze patterns in legal documents and outcomes to forecast potential issues and opportunities")
    
    # Mock predictions
    predictions = [
        {"type": "Contract Risk", "prediction": "15% increase in liability clauses", "confidence": "82%"},
        {"type": "Compliance", "prediction": "New regulations will affect 23% of contracts", "confidence": "76%"},
        {"type": "Case Outcome", "prediction": "Settlement recommended in 67% of cases", "confidence": "89%"}
    ]
    
    for pred in predictions:
        with st.expander(f"{pred['type']}: {pred['prediction']}"):
            st.write(f"**Confidence Level:** {pred['confidence']}")
            st.write("**Recommendation:** Based on historical data analysis...")

if __name__ == "__main__":
    show()
